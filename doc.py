import re
import pandas as pd
from docx import Document
from os import listdir, makedirs, getcwd
from os.path import isfile, join, basename, exists
#склонировал в другую папку и добавил коммит

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def get_docs(doc, docs, names, names_to_xlsx):
    name_to_list = '-'
    doc = Document(doc)
    paragraphs = doc.paragraphs

    for p in paragraphs:
        name = re.findall(r'\w*\s\w\.\w\.', p.text)
        if name:
            name_to_list = name
            for n in name:
                names_to_xlsx.append(n)

    if name_to_list:
        names.append(name_to_list[0])
    docs.append(doc)


def formating(doc):
    NW = [
        'Ключевые слова',
        'Аннотация'
    ]

    paragraphs = doc.paragraphs

    for i in range(5 if len(paragraphs) > 5 else len(paragraphs)):
        p = paragraphs[i]
        for w in NW:
            if re.match(w, p.text):
                delete_paragraph(p)

        name = re.findall(r'\w*\s\w\.\w\.', p.text)
        if name:
            for n in name:
                p.style = doc.styles['Heading 2']
            paragraphs[i + 1].style = doc.styles['Heading 1']


def main():
    # in data
    PATH = getcwd() #нужен комментарий
    FOLDER = 'секции'
    #ещё коммент
    folder = basename(PATH) + ' ' + FOLDER
    if not exists(folder):
        makedirs(folder)
#таки получилось
    files = [f for f in listdir() if isfile(join(f))]
    in_docs = [d for d in files if re.search(r'docx', d)]

    docs = []
    names = []
    names_to_xlsx = []

    # sort files and get one new word
    for doc in in_docs:
        get_docs(doc, docs, names, names_to_xlsx)

    docs_to_sort = dict(zip(names, docs))

    sorted_docs = []
    for s in sorted(docs_to_sort.keys()):
        sorted_docs.append(docs_to_sort[s])

    # formating
    for i, d in enumerate(sorted_docs):
        formating(d)
        d.save(folder + '\\' + folder + str(i) + '.docx')

    # save
    df = pd.DataFrame(sorted(names_to_xlsx), columns=['ФИО'])
    df['Должность'] = ''
    df = df.drop_duplicates()

    writer = pd.ExcelWriter(folder + '\\' + 'names.xlsx', engine='xlsxwriter')
    df.to_excel(writer, sheet_name='Sheet1', index=False)
    writer.save()

    # out_doc.save(folder + '\\' + folder + '.docx')


if __name__ == '__main__':
    main()

